"""
PII Detection & Redaction 服務 — 使用 Microsoft Presidio

功能：
  - 掃描純文字中的 PII（個人可識別資訊）
  - 遮蔽/替換 PII 實體
  - 從檔案提取文字並掃描（支援 PDF、DOCX、TXT、CSV）
  - 自訂台灣特有的 PII Recognizer（身分證、手機、銀行帳號）

使用方式：
  from services.pii_service import pii_service
  result = pii_service.scan_text("我的身分證是A123456789")
  redacted = pii_service.anonymize_text("我的身分證是A123456789")
"""
import logging
import re
from dataclasses import dataclass, field, asdict
from datetime import datetime, timezone
from pathlib import Path
from typing import List, Optional

logger = logging.getLogger(__name__)

# ============================================================
# 資料模型
# ============================================================


@dataclass
class PIIEntity:
    """偵測到的單一 PII 實體"""
    entity_type: str       # e.g. "EMAIL_ADDRESS", "TW_ID"
    text: str              # 偵測到的原始文字
    start: int             # 起始位置
    end: int               # 結束位置
    score: float           # 信心分數 0-1

    def to_dict(self) -> dict:
        return asdict(self)


@dataclass
class PIIScanResult:
    """PII 掃描結果"""
    has_pii: bool = False
    entity_count: int = 0
    entities: List[PIIEntity] = field(default_factory=list)
    entity_types: List[str] = field(default_factory=list)
    scanned_at: str = ""
    text_length: int = 0
    confidence_threshold: float = 0.5

    def to_dict(self) -> dict:
        return {
            "has_pii": self.has_pii,
            "entity_count": self.entity_count,
            "entities": [e.to_dict() for e in self.entities],
            "entity_types": self.entity_types,
            "scanned_at": self.scanned_at,
            "text_length": self.text_length,
            "confidence_threshold": self.confidence_threshold,
        }


# ============================================================
# 台灣特有 PII Recognizer
# ============================================================

# 台灣身分證字號驗證
_TW_ID_PATTERN = re.compile(r'[A-Z][12]\d{8}')

# 台灣身分證字號驗證碼計算用的對照表
_TW_ID_LETTER_MAP = {
    'A': 10, 'B': 11, 'C': 12, 'D': 13, 'E': 14, 'F': 15,
    'G': 16, 'H': 17, 'I': 34, 'J': 18, 'K': 19, 'L': 20,
    'M': 21, 'N': 22, 'O': 35, 'P': 23, 'Q': 24, 'R': 25,
    'S': 26, 'T': 27, 'U': 28, 'V': 29, 'W': 32, 'X': 30,
    'Y': 31, 'Z': 33,
}


def _validate_tw_id(id_str: str) -> bool:
    """驗證台灣身分證字號的驗證碼"""
    if not _TW_ID_PATTERN.match(id_str):
        return False
    letter = id_str[0].upper()
    if letter not in _TW_ID_LETTER_MAP:
        return False
    n = _TW_ID_LETTER_MAP[letter]
    total = (n // 10) + (n % 10) * 9
    weights = [8, 7, 6, 5, 4, 3, 2, 1]
    for i, w in enumerate(weights):
        total += int(id_str[i + 1]) * w
    total += int(id_str[9])
    return total % 10 == 0


# 台灣手機號碼
_TW_MOBILE_PATTERN = re.compile(r'09\d{2}[-\s]?\d{3}[-\s]?\d{3}')

# 台灣市話
_TW_LANDLINE_PATTERN = re.compile(r'0[2-9]\d?[-\s]?\d{3,4}[-\s]?\d{3,4}')

# Email
_EMAIL_PATTERN = re.compile(
    r'[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}'
)

# 信用卡號（4 組 4 位數字，可用空格或連字號分隔）
_CREDIT_CARD_PATTERN = re.compile(
    r'\b(?:\d{4}[-\s]?){3}\d{4}\b'
)

# IP 位址
_IP_PATTERN = re.compile(
    r'\b(?:(?:25[0-5]|2[0-4]\d|[01]?\d\d?)\.){3}(?:25[0-5]|2[0-4]\d|[01]?\d\d?)\b'
)

# 台灣統一編號（8 位數字）
_TW_UBN_PATTERN = re.compile(r'\b\d{8}\b')


def _validate_tw_ubn(ubn_str: str) -> bool:
    """驗證台灣統一編號"""
    if len(ubn_str) != 8 or not ubn_str.isdigit():
        return False
    weights = [1, 2, 1, 2, 1, 2, 4, 1]
    total = 0
    for i, w in enumerate(weights):
        product = int(ubn_str[i]) * w
        total += product // 10 + product % 10
    if total % 5 == 0:
        return True
    # 第 7 碼為 7 時，可能有兩種驗證方式
    if int(ubn_str[6]) == 7 and (total + 1) % 5 == 0:
        return True
    return False


# ============================================================
# PII 服務（使用 Presidio 或 Fallback 正則）
# ============================================================


class PIIService:
    """
    PII 偵測與遮蔽服務

    優先使用 Microsoft Presidio（如果已安裝），
    否則 fallback 到內建的正則表達式引擎。
    """

    def __init__(
        self,
        enabled: bool = True,
        languages: List[str] = None,
        confidence_threshold: float = 0.5,
        redact_mode: str = "replace",
        excluded_entity_types: List[str] = None,
    ):
        self.enabled = enabled
        self.languages = languages or ["en", "zh"]
        self.confidence_threshold = confidence_threshold
        self.redact_mode = redact_mode  # "replace", "mask", "hash"
        # 排除不需要偵測的 entity type（如 DATE_TIME, LOCATION, URL 等）
        self.excluded_entity_types = set(excluded_entity_types or [])

        self._presidio_available = False
        self._analyzer = None
        self._anonymizer = None

        if not enabled:
            logger.info("🔒 PII 服務已停用（PII_ENABLED=false）")
            return

        # 嘗試載入 Presidio
        try:
            from presidio_analyzer import AnalyzerEngine, PatternRecognizer, Pattern
            from presidio_anonymizer import AnonymizerEngine
            from presidio_anonymizer.entities import OperatorConfig

            # 建立 Analyzer
            self._analyzer = AnalyzerEngine()

            # 註冊台灣特有的 Recognizer
            self._register_tw_recognizers(PatternRecognizer, Pattern)

            # 建立 Anonymizer
            self._anonymizer = AnonymizerEngine()

            self._presidio_available = True
            excluded_info = f", 排除: {list(self.excluded_entity_types)}" if self.excluded_entity_types else ""
            logger.info(
                f"🔒 PII 服務已啟用（Presidio 模式）"
                f" — 語言: {self.languages}, 閾值: {self.confidence_threshold}{excluded_info}"
            )
        except ImportError:
            logger.warning(
                "⚠️ Presidio 未安裝，PII 服務將使用內建正則引擎（功能有限）。"
                " 安裝方式: pip install presidio-analyzer presidio-anonymizer"
            )
            logger.info(
                f"🔒 PII 服務已啟用（正則 Fallback 模式）"
                f" — 閾值: {self.confidence_threshold}"
            )

    def _register_tw_recognizers(self, PatternRecognizer, Pattern):
        """註冊台灣特有的 PII Recognizer 到 Presidio"""
        # 台灣身分證字號
        tw_id_recognizer = PatternRecognizer(
            supported_entity="TW_ID",
            name="Taiwan ID Recognizer",
            patterns=[
                Pattern(
                    name="tw_id",
                    regex=r'[A-Z][12]\d{8}',
                    score=0.7,
                ),
            ],
            supported_language="en",  # 正則不分語言
        )
        self._analyzer.registry.add_recognizer(tw_id_recognizer)

        # 台灣手機號碼
        tw_mobile_recognizer = PatternRecognizer(
            supported_entity="TW_PHONE",
            name="Taiwan Mobile Recognizer",
            patterns=[
                Pattern(
                    name="tw_mobile",
                    regex=r'09\d{2}[-\s]?\d{3}[-\s]?\d{3}',
                    score=0.7,
                ),
            ],
            supported_language="en",
        )
        self._analyzer.registry.add_recognizer(tw_mobile_recognizer)

        # 台灣市話
        tw_landline_recognizer = PatternRecognizer(
            supported_entity="TW_PHONE",
            name="Taiwan Landline Recognizer",
            patterns=[
                Pattern(
                    name="tw_landline",
                    regex=r'0[2-9]\d?[-\s]?\d{3,4}[-\s]?\d{3,4}',
                    score=0.5,
                ),
            ],
            supported_language="en",
        )
        self._analyzer.registry.add_recognizer(tw_landline_recognizer)

        # 台灣統一編號
        tw_ubn_recognizer = PatternRecognizer(
            supported_entity="TW_UBN",
            name="Taiwan UBN Recognizer",
            patterns=[
                Pattern(
                    name="tw_ubn",
                    regex=r'\b\d{8}\b',
                    score=0.3,  # 低分，因為 8 位數字太常見
                ),
            ],
            supported_language="en",
        )
        self._analyzer.registry.add_recognizer(tw_ubn_recognizer)

        logger.info("✅ 已註冊台灣特有 PII Recognizer（TW_ID, TW_PHONE, TW_UBN）")

    # ============================================================
    # 核心功能：掃描文字
    # ============================================================

    def scan_text(self, text: str) -> PIIScanResult:
        """
        掃描純文字中的 PII

        Args:
            text: 要掃描的文字

        Returns:
            PIIScanResult 包含偵測到的 PII 實體
        """
        if not self.enabled or not text or not text.strip():
            return PIIScanResult(
                scanned_at=datetime.now(timezone.utc).isoformat(),
                text_length=len(text) if text else 0,
                confidence_threshold=self.confidence_threshold,
            )

        if self._presidio_available:
            return self._scan_with_presidio(text)
        else:
            return self._scan_with_regex(text)

    def _scan_with_presidio(self, text: str) -> PIIScanResult:
        """使用 Presidio 掃描"""
        try:
            # 對每個語言執行分析
            all_results = []
            for lang in self.languages:
                try:
                    results = self._analyzer.analyze(
                        text=text,
                        language=lang,
                        score_threshold=self.confidence_threshold,
                    )
                    all_results.extend(results)
                except Exception as e:
                    logger.debug(f"Presidio 分析語言 {lang} 失敗: {e}")

            # 去重（同一位置可能被多語言重複偵測）
            seen = set()
            unique_results = []
            for r in all_results:
                key = (r.start, r.end, r.entity_type)
                if key not in seen:
                    seen.add(key)
                    unique_results.append(r)

            # 排除不需要偵測的 entity type（如 DATE_TIME, LOCATION, URL 等）
            if self.excluded_entity_types:
                before_count = len(unique_results)
                unique_results = [
                    r for r in unique_results
                    if r.entity_type not in self.excluded_entity_types
                ]
                excluded_count = before_count - len(unique_results)
                if excluded_count > 0:
                    logger.debug(f"已排除 {excluded_count} 個不需偵測的 entity type")

            # 對台灣身分證做額外驗證
            validated_results = []
            for r in unique_results:
                if r.entity_type == "TW_ID":
                    id_text = text[r.start:r.end]
                    if _validate_tw_id(id_text):
                        r.score = max(r.score, 0.95)  # 驗證通過，提高信心
                        validated_results.append(r)
                    # 驗證不通過就跳過
                elif r.entity_type == "TW_UBN":
                    ubn_text = text[r.start:r.end]
                    if _validate_tw_ubn(ubn_text):
                        r.score = max(r.score, 0.8)
                        validated_results.append(r)
                    # 8 位數字但不是統一編號就跳過
                else:
                    validated_results.append(r)

            entities = [
                PIIEntity(
                    entity_type=r.entity_type,
                    text=text[r.start:r.end],
                    start=r.start,
                    end=r.end,
                    score=round(r.score, 3),
                )
                for r in validated_results
            ]

            entity_types = list(set(e.entity_type for e in entities))

            return PIIScanResult(
                has_pii=len(entities) > 0,
                entity_count=len(entities),
                entities=entities,
                entity_types=entity_types,
                scanned_at=datetime.now(timezone.utc).isoformat(),
                text_length=len(text),
                confidence_threshold=self.confidence_threshold,
            )

        except Exception as e:
            logger.error(f"❌ Presidio 掃描失敗: {e}")
            # Fallback 到正則
            return self._scan_with_regex(text)

    def _scan_with_regex(self, text: str) -> PIIScanResult:
        """使用內建正則表達式掃描（Presidio 不可用時的 fallback）"""
        entities = []

        # 台灣身分證
        for m in _TW_ID_PATTERN.finditer(text):
            if _validate_tw_id(m.group()):
                entities.append(PIIEntity(
                    entity_type="TW_ID",
                    text=m.group(),
                    start=m.start(),
                    end=m.end(),
                    score=0.95,
                ))

        # Email
        for m in _EMAIL_PATTERN.finditer(text):
            entities.append(PIIEntity(
                entity_type="EMAIL_ADDRESS",
                text=m.group(),
                start=m.start(),
                end=m.end(),
                score=0.95,
            ))

        # 台灣手機
        for m in _TW_MOBILE_PATTERN.finditer(text):
            entities.append(PIIEntity(
                entity_type="TW_PHONE",
                text=m.group(),
                start=m.start(),
                end=m.end(),
                score=0.85,
            ))

        # 台灣市話
        for m in _TW_LANDLINE_PATTERN.finditer(text):
            # 避免與手機號碼重疊
            overlap = False
            for e in entities:
                if e.entity_type == "TW_PHONE" and (
                    (m.start() >= e.start and m.start() < e.end) or
                    (m.end() > e.start and m.end() <= e.end)
                ):
                    overlap = True
                    break
            if not overlap:
                entities.append(PIIEntity(
                    entity_type="TW_PHONE",
                    text=m.group(),
                    start=m.start(),
                    end=m.end(),
                    score=0.6,
                ))

        # 信用卡號
        for m in _CREDIT_CARD_PATTERN.finditer(text):
            # 簡單的 Luhn 檢查可以加在這裡
            entities.append(PIIEntity(
                entity_type="CREDIT_CARD",
                text=m.group(),
                start=m.start(),
                end=m.end(),
                score=0.7,
            ))

        # IP 位址
        for m in _IP_PATTERN.finditer(text):
            entities.append(PIIEntity(
                entity_type="IP_ADDRESS",
                text=m.group(),
                start=m.start(),
                end=m.end(),
                score=0.7,
            ))

        # 台灣統一編號
        for m in _TW_UBN_PATTERN.finditer(text):
            if _validate_tw_ubn(m.group()):
                # 確認不與其他已偵測的實體重疊
                overlap = False
                for e in entities:
                    if (m.start() >= e.start and m.start() < e.end) or \
                       (m.end() > e.start and m.end() <= e.end):
                        overlap = True
                        break
                if not overlap:
                    entities.append(PIIEntity(
                        entity_type="TW_UBN",
                        text=m.group(),
                        start=m.start(),
                        end=m.end(),
                        score=0.8,
                    ))

        # 過濾低於閾值的結果
        entities = [e for e in entities if e.score >= self.confidence_threshold]

        # 按位置排序
        entities.sort(key=lambda e: e.start)

        entity_types = list(set(e.entity_type for e in entities))

        return PIIScanResult(
            has_pii=len(entities) > 0,
            entity_count=len(entities),
            entities=entities,
            entity_types=entity_types,
            scanned_at=datetime.now(timezone.utc).isoformat(),
            text_length=len(text),
            confidence_threshold=self.confidence_threshold,
        )

    # ============================================================
    # 核心功能：遮蔽文字
    # ============================================================

    def anonymize_text(self, text: str) -> str:
        """
        遮蔽文字中的 PII

        Args:
            text: 要遮蔽的文字

        Returns:
            脫敏後的文字
        """
        if not self.enabled or not text or not text.strip():
            return text

        if self._presidio_available:
            return self._anonymize_with_presidio(text)
        else:
            return self._anonymize_with_regex(text)

    def _anonymize_with_presidio(self, text: str) -> str:
        """使用 Presidio Anonymizer 遮蔽"""
        try:
            from presidio_anonymizer.entities import OperatorConfig

            # 先分析
            all_results = []
            for lang in self.languages:
                try:
                    results = self._analyzer.analyze(
                        text=text,
                        language=lang,
                        score_threshold=self.confidence_threshold,
                    )
                    all_results.extend(results)
                except Exception:
                    pass

            if not all_results:
                return text

            # 去重
            seen = set()
            unique_results = []
            for r in all_results:
                key = (r.start, r.end, r.entity_type)
                if key not in seen:
                    seen.add(key)
                    unique_results.append(r)

            # 排除不需要偵測的 entity type
            if self.excluded_entity_types:
                unique_results = [
                    r for r in unique_results
                    if r.entity_type not in self.excluded_entity_types
                ]

            # 台灣身分證額外驗證
            validated_results = []
            for r in unique_results:
                if r.entity_type == "TW_ID":
                    id_text = text[r.start:r.end]
                    if _validate_tw_id(id_text):
                        validated_results.append(r)
                elif r.entity_type == "TW_UBN":
                    ubn_text = text[r.start:r.end]
                    if _validate_tw_ubn(ubn_text):
                        validated_results.append(r)
                else:
                    validated_results.append(r)

            if not validated_results:
                return text

            # 根據 redact_mode 設定遮蔽方式
            if self.redact_mode == "mask":
                operators = {
                    "DEFAULT": OperatorConfig("mask", {"chars_to_mask": 100, "masking_char": "*", "from_end": False}),
                }
            elif self.redact_mode == "hash":
                operators = {
                    "DEFAULT": OperatorConfig("hash", {"hash_type": "sha256"}),
                }
            else:  # "replace" (預設)
                operators = {
                    "DEFAULT": OperatorConfig("replace", {"new_value": "<REDACTED>"}),
                }
                # 為每種類型設定特定的替換文字
                for r in validated_results:
                    if r.entity_type not in operators:
                        operators[r.entity_type] = OperatorConfig(
                            "replace",
                            {"new_value": f"<{r.entity_type}>"},
                        )

            result = self._anonymizer.anonymize(
                text=text,
                analyzer_results=validated_results,
                operators=operators,
            )
            return result.text

        except Exception as e:
            logger.error(f"❌ Presidio 遮蔽失敗: {e}")
            return self._anonymize_with_regex(text)

    def _anonymize_with_regex(self, text: str) -> str:
        """使用正則表達式遮蔽（fallback）"""
        scan_result = self._scan_with_regex(text)
        if not scan_result.has_pii:
            return text

        # 從後往前替換（避免位置偏移）
        result = text
        for entity in sorted(scan_result.entities, key=lambda e: e.start, reverse=True):
            if self.redact_mode == "mask":
                replacement = "*" * len(entity.text)
            elif self.redact_mode == "hash":
                import hashlib
                replacement = hashlib.sha256(entity.text.encode()).hexdigest()[:16]
            else:  # "replace"
                replacement = f"<{entity.entity_type}>"
            result = result[:entity.start] + replacement + result[entity.end:]

        return result

    # ============================================================
    # 檔案掃描
    # ============================================================

    async def scan_file(self, file_path: Path) -> PIIScanResult:
        """
        從檔案提取文字並掃描 PII

        支援格式：PDF、DOCX、TXT、CSV

        Args:
            file_path: 檔案路徑

        Returns:
            PIIScanResult
        """
        if not self.enabled:
            return PIIScanResult(
                scanned_at=datetime.now(timezone.utc).isoformat(),
                confidence_threshold=self.confidence_threshold,
            )

        if not file_path.exists():
            logger.warning(f"⚠️ 檔案不存在: {file_path}")
            return PIIScanResult(
                scanned_at=datetime.now(timezone.utc).isoformat(),
                confidence_threshold=self.confidence_threshold,
            )

        ext = file_path.suffix.lower()
        text = ""

        try:
            if ext == ".txt" or ext == ".csv":
                text = self._extract_text_from_txt(file_path)
            elif ext == ".pdf":
                text = self._extract_text_from_pdf(file_path)
            elif ext in (".docx", ".doc"):
                text = self._extract_text_from_docx(file_path)
            else:
                logger.info(f"ℹ️ 不支援的檔案格式，跳過 PII 掃描: {ext}")
                return PIIScanResult(
                    scanned_at=datetime.now(timezone.utc).isoformat(),
                    confidence_threshold=self.confidence_threshold,
                )
        except Exception as e:
            logger.error(f"❌ 檔案文字提取失敗 ({file_path}): {e}")
            return PIIScanResult(
                scanned_at=datetime.now(timezone.utc).isoformat(),
                confidence_threshold=self.confidence_threshold,
            )

        if not text.strip():
            return PIIScanResult(
                scanned_at=datetime.now(timezone.utc).isoformat(),
                text_length=0,
                confidence_threshold=self.confidence_threshold,
            )

        # 限制掃描文字長度（避免超大檔案耗時過長）
        max_chars = 100_000  # 10 萬字
        if len(text) > max_chars:
            logger.info(f"ℹ️ 文字過長 ({len(text)} chars)，僅掃描前 {max_chars} 字")
            text = text[:max_chars]

        return self.scan_text(text)

    def _extract_text_from_txt(self, file_path: Path) -> str:
        """從 TXT/CSV 提取文字"""
        encodings = ["utf-8", "big5", "gb2312", "latin-1"]
        for enc in encodings:
            try:
                return file_path.read_text(encoding=enc)
            except (UnicodeDecodeError, UnicodeError):
                continue
        return ""

    def _extract_text_from_pdf(self, file_path: Path) -> str:
        """從 PDF 提取文字"""
        try:
            import pdfplumber
            text_parts = []
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
            return "\n".join(text_parts)
        except ImportError:
            logger.warning("⚠️ pdfplumber 未安裝，無法提取 PDF 文字。安裝: pip install pdfplumber")
            return ""
        except Exception as e:
            logger.error(f"❌ PDF 文字提取失敗: {e}")
            return ""

    def _extract_text_from_docx(self, file_path: Path) -> str:
        """從 DOCX 提取文字"""
        try:
            from docx import Document
            doc = Document(str(file_path))
            text_parts = [p.text for p in doc.paragraphs if p.text]
            # 也提取表格中的文字
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text:
                            text_parts.append(cell.text)
            return "\n".join(text_parts)
        except ImportError:
            logger.warning("⚠️ python-docx 未安裝，無法提取 DOCX 文字。安裝: pip install python-docx")
            return ""
        except Exception as e:
            logger.error(f"❌ DOCX 文字提取失敗: {e}")
            return ""

    # ============================================================
    # 工具方法
    # ============================================================

    def get_status(self) -> dict:
        """取得 PII 服務狀態"""
        return {
            "enabled": self.enabled,
            "engine": "presidio" if self._presidio_available else "regex_fallback",
            "languages": self.languages,
            "confidence_threshold": self.confidence_threshold,
            "redact_mode": self.redact_mode,
            "excluded_entity_types": list(self.excluded_entity_types),
        }


# ============================================================
# 模組級別的 singleton 實例（延遲初始化）
# ============================================================

_pii_service: Optional[PIIService] = None


def get_pii_service() -> PIIService:
    """取得 PII 服務的 singleton 實例"""
    global _pii_service
    if _pii_service is None:
        _pii_service = _create_pii_service()
    return _pii_service


def _create_pii_service() -> PIIService:
    """根據 config 建立 PII 服務"""
    try:
        from config import settings
        return PIIService(
            enabled=settings.PII_ENABLED,
            languages=settings.PII_LANGUAGES,
            confidence_threshold=settings.PII_CONFIDENCE_THRESHOLD,
            redact_mode=settings.PII_REDACT_MODE,
            excluded_entity_types=settings.PII_EXCLUDED_ENTITIES,
        )
    except Exception as e:
        logger.warning(f"⚠️ 無法從 config 載入 PII 設定，使用預設值: {e}")
        return PIIService(enabled=False)


# 方便直接 import 使用的 alias
pii_service = property(lambda self: get_pii_service())
